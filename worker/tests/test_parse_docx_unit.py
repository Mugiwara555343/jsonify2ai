# worker/tests/test_parse_docx_unit.py
import pytest

try:
    from docx import Document  # type: ignore

    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False

from app.services.parse_docx import extract_text_from_docx

pytestmark = pytest.mark.skipif(
    not DOCX_AVAILABLE,
    reason="python-docx not installed; install with 'pip install python-docx' to run this test",
)


def test_extract_text_from_docx(tmp_path):
    p = tmp_path / "s.docx"
    doc = Document()
    doc.add_paragraph("Hello world")
    row = doc.add_table(rows=1, cols=2).rows[0]
    row.cells[0].text, row.cells[1].text = "A", "B"
    doc.save(p)

    txt = extract_text_from_docx(str(p))
    assert "Hello world" in txt
    assert "A" in txt and "B" in txt
