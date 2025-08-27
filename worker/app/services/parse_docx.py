# worker/app/services/parse_docx.py
def extract_text_from_docx(path: str) -> str:
    try:
        from docx import Document
    except Exception as e:
        raise RuntimeError(
            "python-docx is required to parse .docx files. Install with: pip install python-docx"
        ) from e
    doc = Document(path)
    lines = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            lines.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(lines)
